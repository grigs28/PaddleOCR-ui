"""Office 文档处理：docx/xlsx 文本提取 + LibreOffice 转 PDF（可选）+ DWG DXF 文字提取"""

import asyncio
import io
import json
import logging
import os
import re
import shutil

from backend.config import get_settings

logger = logging.getLogger(__name__)

# 检测 LibreOffice 是否可用
_libreoffice_path: str | None = None

# 检测 cad2x 是否可用
_cad2x_path: str | None = None


def _find_libreoffice() -> str | None:
    for cmd in ("libreoffice", "soffice", "/usr/bin/libreoffice", "/usr/bin/soffice"):
        if shutil.which(cmd):
            return cmd
    return None


def is_libreoffice_available() -> bool:
    global _libreoffice_path
    if _libreoffice_path is None:
        _libreoffice_path = _find_libreoffice()
    return _libreoffice_path is not None


def _find_cad2x() -> str | None:
    """查找 cad2x 二进制"""
    # 项目 bin 目录优先
    project_bin = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "bin", "cad2x")
    if os.path.isfile(project_bin) and os.access(project_bin, os.X_OK):
        return project_bin
    # 系统 PATH
    path = shutil.which("cad2x")
    if path:
        return path
    return None


def is_cad2x_available() -> bool:
    global _cad2x_path
    if _cad2x_path is None:
        _cad2x_path = _find_cad2x()
    return _cad2x_path is not None


async def convert_dwg_to_pdf(input_path: str, output_dir: str) -> str | None:
    """DWG/DXF 转 PDF。优先用 ACAD 服务（按图框分页），回退到 cad2x。

    ACAD 服务可能返回多个 PDF（每个图框一个），合并为一个 PDF 返回。
    返回最终 PDF 路径，失败返回 None。
    """
    # 优先尝试 ACAD 服务
    pdf_path = await _convert_dwg_via_acad(input_path, output_dir)
    if pdf_path:
        return pdf_path

    # 回退到 cad2x
    logger.info("ACAD 服务不可用或失败，回退到 cad2x")
    return await _convert_dwg_via_cad2x(input_path, output_dir)


async def _acad_request(client, method: str, path: str, **kwargs):
    """发送带认证的 ACAD API 请求"""
    import aiohttp
    settings = get_settings()
    url = settings.acad_service_url.rstrip("/") + path
    headers = kwargs.pop("headers", {})
    if settings.acad_service_apikey:
        headers["x-api-key"] = settings.acad_service_apikey
    return await client.request(method, url, headers=headers, **kwargs)


async def _poll_acad_task(client, task_path: str, task_id: str) -> dict | None:
    """轮询 ACAD 任务状态，返回完成的 task 数据或 None"""
    settings = get_settings()
    timeout = settings.libreoffice_timeout
    deadline = asyncio.get_event_loop().time() + timeout

    while asyncio.get_event_loop().time() < deadline:
        try:
            resp = await _acad_request(client, "GET", f"{task_path}/{task_id}")
            if resp.status == 200:
                data = await resp.json()
                if data.get("status") == "done":
                    return data
                if data.get("status") == "failed":
                    errors = [f["error"] for f in data.get("files", []) if f.get("error")]
                    logger.warning(f"ACAD 任务失败: {errors}")
                    return None
            await asyncio.sleep(5)
        except Exception as e:
            logger.warning(f"ACAD 轮询异常: {e}")
            await asyncio.sleep(5)
    logger.warning(f"ACAD 任务超时 ({timeout}s)")
    return None


async def _download_acad_result(task_id: str, download_path: str, output_dir: str) -> list[str]:
    """下载 ACAD 任务结果 ZIP 并解压到 output_dir，返回解压后的文件路径列表"""
    import aiohttp
    import zipfile

    async with aiohttp.ClientSession() as client:
        resp = await _acad_request(client, "GET", f"{download_path}/{task_id}")
        if resp.status != 200:
            logger.warning(f"ACAD 下载失败: {resp.status}")
            return []
        zip_bytes = await resp.read()

    if not zip_bytes:
        return []

    temp_extract = os.path.join(output_dir, f"_acad_temp_{task_id[:8]}")
    os.makedirs(temp_extract, exist_ok=True)
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            zf.extractall(temp_extract)
        result_files = []
        for root, dirs, files in os.walk(temp_extract):
            for f in files:
                result_files.append(os.path.join(root, f))
        return result_files
    except Exception as e:
        logger.warning(f"ACAD 解压失败: {e}")
        shutil.rmtree(temp_extract, ignore_errors=True)
        return []


async def _convert_dwg_via_acad(input_path: str, output_dir: str) -> str | None:
    """通过 ACADxPDF 服务将 DWG 转 PDF+DXF（异步轮询模式）。
    返回合并后的 PDF 路径，DXF 保存在同目录。
    """
    import aiohttp

    settings = get_settings()

    # 先检查服务是否可用
    try:
        async with aiohttp.ClientSession() as client:
            resp = await _acad_request(client, "GET", "/health")
            if resp.status != 200:
                return None
    except Exception:
        return None

    # 提交转换任务
    try:
        async with aiohttp.ClientSession() as client:
            with open(input_path, "rb") as f:
                data = aiohttp.FormData()
                data.add_field("files", f, filename=os.path.basename(input_path))
                resp = await _acad_request(client, "POST", "/convert", data=data,
                                           timeout=aiohttp.ClientTimeout(total=30))
                if resp.status != 200:
                    text = await resp.text()
                    logger.warning(f"ACAD 提交失败 {resp.status}: {text[:200]}")
                    return None
                result = await resp.json()
                task_id = result.get("task_id")
                if not task_id:
                    return None
    except Exception as e:
        logger.warning(f"ACAD 提交异常: {e}")
        return None

    # 轮询任务状态
    async with aiohttp.ClientSession() as client:
        task_data = await _poll_acad_task(client, "/tasks", task_id)
    if not task_data:
        return None

    # 下载结果
    extracted = await _download_acad_result(task_id, "/download", output_dir)
    if not extracted:
        return None

    # 分类文件
    pdf_files = []
    dxf_files = []
    for f in extracted:
        lower = f.lower()
        if lower.endswith(".pdf"):
            pdf_files.append(f)
        elif lower.endswith(".dxf"):
            dxf_files.append(f)

    # 移动 DXF 到 output_dir
    for dxf in dxf_files:
        dxf_dest = os.path.join(output_dir, os.path.basename(dxf))
        if not os.path.exists(dxf_dest):
            shutil.move(dxf, dxf_dest)
            logger.info(f"DXF 已保存: {dxf_dest}")

    if not pdf_files:
        shutil.rmtree(os.path.dirname(extracted[0]), ignore_errors=True)
        return None

    # 移动 PDF 到 output_dir
    saved_pdfs = []
    for pdf in pdf_files:
        dest = os.path.join(output_dir, os.path.basename(pdf))
        if not os.path.exists(dest):
            shutil.copy2(pdf, dest)
        saved_pdfs.append(dest)

    base = os.path.splitext(os.path.basename(input_path))[0]

    if len(saved_pdfs) == 1:
        final_path = saved_pdfs[0]
        shutil.rmtree(os.path.dirname(extracted[0]), ignore_errors=True)
        logger.info(f"ACAD 转换成功（单页）: {final_path}")
        return final_path

    # 多页合并
    final_path = os.path.join(output_dir, f"{base}.pdf")
    await _merge_pdfs(saved_pdfs, final_path)
    shutil.rmtree(os.path.dirname(extracted[0]), ignore_errors=True)
    logger.info(f"ACAD 转换成功（{len(pdf_files)} 页合并）: {final_path}")
    return final_path


async def convert_pdf_to_dwg(input_path: str, output_dir: str) -> str | None:
    """通过 ACADxPDF 服务将 PDF 转 DWG（异步轮询模式）。
    返回 DWG 文件路径，失败返回 None。
    """
    import aiohttp

    # 检查服务
    try:
        async with aiohttp.ClientSession() as client:
            resp = await _acad_request(client, "GET", "/health")
            if resp.status != 200:
                return None
    except Exception:
        return None

    # 提交 PDF→DWG 任务
    try:
        async with aiohttp.ClientSession() as client:
            with open(input_path, "rb") as f:
                data = aiohttp.FormData()
                data.add_field("files", f, filename=os.path.basename(input_path))
                resp = await _acad_request(client, "POST", "/convert-pdf", data=data,
                                           timeout=aiohttp.ClientTimeout(total=30))
                if resp.status != 200:
                    text = await resp.text()
                    logger.warning(f"ACAD PDF→DWG 提交失败 {resp.status}: {text[:200]}")
                    return None
                result = await resp.json()
                task_id = result.get("task_id")
                if not task_id:
                    return None
    except Exception as e:
        logger.warning(f"ACAD PDF→DWG 提交异常: {e}")
        return None

    # 轮询
    async with aiohttp.ClientSession() as client:
        task_data = await _poll_acad_task(client, "/pdf-task", task_id)
    if not task_data:
        return None

    # 下载结果
    extracted = await _download_acad_result(task_id, "/download-pdf-zip", output_dir)
    if not extracted:
        return None

    # 找 DWG 文件
    dwg_files = [f for f in extracted if f.lower().endswith(".dwg")]
    if not dwg_files:
        logger.warning("PDF→DWG 结果中无 DWG 文件")
        shutil.rmtree(os.path.dirname(extracted[0]), ignore_errors=True)
        return None

    # 移动 DWG 到 output_dir
    dwg_dest = os.path.join(output_dir, os.path.basename(dwg_files[0]))
    shutil.copy2(dwg_files[0], dwg_dest)
    shutil.rmtree(os.path.dirname(extracted[0]), ignore_errors=True)
    logger.info(f"PDF→DWG 转换成功: {dwg_dest}")
    return dwg_dest


async def _merge_pdfs(pdf_paths: list[str], output_path: str) -> None:
    """合并多个 PDF 文件为一个（使用系统 pdfunite）"""
    proc = await asyncio.create_subprocess_exec(
        "pdfunite", *pdf_paths, output_path,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    stdout, stderr = await proc.communicate()
    if proc.returncode != 0:
        raise Exception(f"pdfunite 失败: {stderr.decode()[:200]}")


async def _convert_dwg_via_cad2x(input_path: str, output_dir: str) -> str | None:
    """用 cad2x 将 DWG/DXF 转为 PDF（单页）。返回 PDF 路径，失败返回 None。"""
    if not is_cad2x_available():
        logger.error("cad2x 不可用")
        return None

    base = os.path.splitext(os.path.basename(input_path))[0]
    pdf_path = os.path.join(output_dir, f"{base}.pdf")

    cmd = [
        _cad2x_path,
        "-o", pdf_path,
        input_path,
        "-ac",          # 自动方向 + 居中
        "-e", "ANSI_936",   # 简体中文编码
        "-f", "simsun",     # 宋体
    ]
    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=get_settings().libreoffice_timeout)
        if proc.returncode == 0 and os.path.exists(pdf_path):
            logger.info(f"cad2x 转换成功: {pdf_path}")
            return pdf_path
        logger.warning(f"cad2x 转换失败: rc={proc.returncode}, stderr={stderr.decode()[:200]}")
    except asyncio.TimeoutError:
        logger.warning("cad2x 转换超时")
        try:
            proc.kill()
        except Exception:
            pass
    except Exception as e:
        logger.warning(f"cad2x 异常: {e}")
    return None


async def convert_to_pdf(input_path: str, output_dir: str) -> str | None:
    """
    用 LibreOffice headless 将 Office 文件转 PDF。
    返回 PDF 路径，失败返回 None。
    """
    if not is_libreoffice_available():
        return None

    cmd = [
        _libreoffice_path,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        input_path,
    ]
    try:
        import tempfile
        # 用临时目录做 user profile，避免多实例锁冲突
        profile_dir = tempfile.mkdtemp(prefix="lo_profile_")
        full_cmd = [_libreoffice_path, f"-env:UserInstallation=file://{profile_dir}"] + cmd[1:]
        proc = await asyncio.create_subprocess_exec(
            *full_cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=get_settings().libreoffice_timeout)
        if proc.returncode == 0:
            base = os.path.splitext(os.path.basename(input_path))[0]
            pdf_path = os.path.join(output_dir, f"{base}.pdf")
            if os.path.exists(pdf_path):
                return pdf_path
        logger.warning(f"LibreOffice 转换失败: rc={proc.returncode}, stderr={stderr.decode()[:200]}")
    except asyncio.TimeoutError:
        logger.warning("LibreOffice 转换超时(3600s)")
        try:
            proc.kill()
        except Exception:
            pass
    except Exception as e:
        logger.warning(f"LibreOffice 异常: {e}")
    finally:
        # 清理临时 profile
        try:
            import shutil
            shutil.rmtree(profile_dir, ignore_errors=True)
        except Exception:
            pass
    return None


def is_legacy_office(filename: str) -> bool:
    """doc/xls 旧格式，python 库不支持，必须走 LibreOffice"""
    ext = os.path.splitext(filename)[1].lstrip(".").lower()
    return ext in ("doc", "xls")


def extract_docx_text(file_path: str) -> dict:
    """提取 docx 文本，返回 {markdown, pages, structured}"""
    from docx import Document

    doc = Document(file_path)
    parts = []
    structured_pages = []
    blocks = []
    block_id = 0

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        # 根据样式判断类型
        if "Heading 1" in style:
            label = "doc_title"
            md = f"# {text}"
        elif "Heading" in style:
            label = "paragraph_title"
            level = style.replace("Heading ", "").strip()
            try:
                lvl = int(level)
            except ValueError:
                lvl = 2
            md = f"{'#' * min(lvl, 6)} {text}"
        else:
            label = "text"
            md = text

        parts.append(md)
        blocks.append({
            "id": block_id,
            "global_id": block_id,
            "type": label,
            "content": text,
            "bbox": None,
            "order": None,
        })
        block_id += 1

    # 提取表格
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if rows_data:
            # markdown 表格
            header = "| " + " | ".join(rows_data[0]) + " |"
            sep = "| " + " | ".join(["---"] * len(rows_data[0])) + " |"
            body = "\n".join("| " + " | ".join(r) + " |" for r in rows_data[1:])
            table_md = f"{header}\n{sep}\n{body}"
            parts.append(table_md)
            blocks.append({
                "id": block_id,
                "global_id": block_id,
                "type": "table",
                "content": json.dumps(rows_data, ensure_ascii=False),
                "bbox": None,
                "order": None,
            })
            block_id += 1

    markdown = "\n\n".join(parts)
    structured_pages.append({
        "page": 1,
        "width": None,
        "height": None,
        "blocks": blocks,
    })

    return {
        "markdown": markdown,
        "pages": 1,
        "structured": structured_pages,
    }


def extract_xlsx_text(file_path: str) -> dict:
    """提取 xlsx 文本，返回 {markdown, pages, structured}"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    structured_pages = []
    global_block_id = 0

    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        blocks = []
        rows_data = []

        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            # 跳过全空行
            if not any(cells):
                continue
            rows_data.append(cells)

        if not rows_data:
            continue

        # markdown 表格
        header = "| " + " | ".join(rows_data[0]) + " |"
        sep = "| " + " | ".join(["---"] * len(rows_data[0])) + " |"
        body = "\n".join("| " + " | ".join(r) + " |" for r in rows_data[1:])
        table_md = f"## {sheet_name}\n\n{header}\n{sep}\n{body}"

        parts.append(table_md)
        blocks.append({
            "id": 0,
            "global_id": global_block_id,
            "type": "table",
            "content": json.dumps(rows_data, ensure_ascii=False),
            "bbox": None,
            "order": None,
        })
        global_block_id += 1

        structured_pages.append({
            "page": sheet_idx + 1,
            "width": None,
            "height": None,
            "blocks": blocks,
        })

    wb.close()
    markdown = "\n\n".join(parts)

    return {
        "markdown": markdown,
        "pages": len(structured_pages),
        "structured": structured_pages,
    }


def extract_dxf_text(dxf_path: str) -> dict:
    """从 DXF 文件提取文字内容，生成 Markdown。
    按 Y 坐标从上到下排列，大字号的作为标题，同行文字自动拼接。

    Returns:
        {"markdown": str, "pages": 1, "structured": [...], "images": {}}
    """
    import ezdxf

    doc = ezdxf.readfile(dxf_path)
    msp = doc.modelspace()

    # 收集所有文字实体
    raw_texts = []
    for e in msp.query('TEXT MTEXT'):
        try:
            if hasattr(e, 'plain_text'):
                content = e.plain_text()
            else:
                content = e.dxf.text
        except Exception:
            continue

        content = content.strip()
        if not content:
            continue
        # 清理 DXF 文字中的控制字符
        content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', content)

        # 获取位置和字号
        try:
            insert = e.dxf.insert
            x = float(insert[0]) if hasattr(insert, '__getitem__') else 0.0
            y = float(insert[1]) if len(insert) > 1 else 0.0
        except Exception:
            x, y = 0.0, 0.0

        try:
            height = float(e.dxf.height)
        except Exception:
            height = 0.0

        raw_texts.append({
            'x': x,
            'y': y,
            'content': content,
            'height': height,
        })

    if not raw_texts:
        return {
            "markdown": "",
            "pages": 1,
            "structured": [],
            "images": {},
        }

    # 按 Y 倒序（从上到下），Y 相同按 X 从左到右
    raw_texts.sort(key=lambda t: (-t['y'], t['x']))

    # 将 Y 坐标接近的文字归为同一行（容差 = 字号的 0.5 倍）
    lines = []
    current_line = [raw_texts[0]]
    for t in raw_texts[1:]:
        prev = current_line[0]
        tolerance = max(prev['height'], t['height']) * 0.5
        if abs(t['y'] - prev['y']) <= tolerance:
            current_line.append(t)
        else:
            lines.append(current_line)
            current_line = [t]
    lines.append(current_line)

    # 生成 Markdown，根据字号判断标题级别
    md_parts = []
    blocks = []
    block_id = 0

    # 统计字号分布，找标题阈值
    heights = [t['height'] for t in raw_texts if t['height'] > 0]
    avg_height = sum(heights) / len(heights) if heights else 300

    for line_texts in lines:
        # 拼接同行文字
        line_texts.sort(key=lambda t: t['x'])  # 按 X 从左到右
        text = ' '.join(t['content'] for t in line_texts)
        max_h = max(t['height'] for t in line_texts)

        if not text.strip():
            continue

        # 根据字号判断标题
        if max_h > avg_height * 2.5:
            md_parts.append(f"# {text}")
            label = "title"
        elif max_h > avg_height * 1.8:
            md_parts.append(f"## {text}")
            label = "section_title"
        elif max_h > avg_height * 1.3:
            md_parts.append(f"### {text}")
            label = "subsection_title"
        else:
            md_parts.append(text)
            label = "text"

        blocks.append({
            "id": block_id,
            "global_id": block_id,
            "type": label,
            "content": text,
            "bbox": None,
            "order": None,
        })
        block_id += 1

    markdown = "\n\n".join(md_parts)

    return {
        "markdown": markdown,
        "pages": 1,
        "structured": [{
            "page": 1,
            "width": None,
            "height": None,
            "blocks": blocks,
        }],
        "images": {},
    }
