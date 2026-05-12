import re
import os
import tempfile
from docx import Document
from docx.shared import Pt

# Unicode 上标/下标映射
_SUP_MAP = {str(i): chr(0x2070 + i) for i in range(10)}
_SUP_MAP.update({'n': 'ⁿ', '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾'})
_SUB_MAP = {str(i): chr(0x2080 + i) for i in range(10)}
_SUB_MAP.update({'+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎'})

_LATEX_SYMBOLS = {
    r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
    r'\pm': '±', r'\times': '×', r'\div': '÷',
    r'\approx': '≈', r'\sim': '∼', r'\infty': '∞',
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
    r'\delta': 'δ', r'\theta': 'θ', r'\pi': 'π',
    r'\mu': 'μ', r'\sigma': 'σ', r'\omega': 'ω',
    r'\partial': '∂',
    "&#x27;": "'", "&#39;": "'",
}


def _clean_xml_text(text: str) -> str:
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)


def _latex_to_unicode(match: re.Match) -> str:
    """将 $...$ LaTeX 片段转为 Unicode 上标/下标/符号"""
    tex = match.group(1).strip()
    result = tex

    def _sup(m):
        return ''.join(_SUP_MAP.get(c, c) for c in m.group(1))
    result = re.sub(r'\^{(.+?)}', _sup, result)

    def _sub(m):
        return ''.join(_SUB_MAP.get(c, c) for c in m.group(1))
    result = re.sub(r'_{(.+?)}', _sub, result)

    for k, v in _LATEX_SYMBOLS.items():
        result = result.replace(k, v)

    return result


def _preprocess_latex(text: str) -> str:
    """预处理 LaTeX $...$ 为 Unicode"""
    return re.sub(r'\$\s*(.+?)\s*\$', _latex_to_unicode, text)


class ExportService:
    @staticmethod
    def md_to_txt(md_text: str) -> str:
        """Markdown 转纯文本"""
        text = md_text
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        return text.strip()

    @staticmethod
    def md_to_docx(md_text: str) -> bytes:
        """Markdown 转 DOCX（支持 HTML 表格 + 标题 + LaTeX 上标/下标）

        流程：预处理 LaTeX → markdown-it 渲染 HTML → pypandoc 转 DOCX
        """
        try:
            import pypandoc
            from markdown_it import MarkdownIt

            # 1. LaTeX → Unicode
            processed = _preprocess_latex(md_text)

            # 2. markdown-it 渲染为 HTML（启用 HTML 内嵌和表格）
            mi = MarkdownIt('commonmark', {'html': True}).enable('table')
            html_body = mi.render(processed)

            # 3. Pandoc HTML → DOCX
            html = '<html><body>' + html_body + '</body></html>'
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
            pypandoc.convert_text(
                html, 'docx', format='html',
                outputfile=tmp_path,
                extra_args=['--wrap=none'],
            )
            with open(tmp_path, 'rb') as f:
                data = f.read()
            os.unlink(tmp_path)
            return data
        except ImportError:
            return _md_to_docx_simple(md_text)


def _md_to_docx_simple(md_text: str) -> bytes:
    """python-docx 简易转换（回退方案）"""
    import io
    doc = Document()
    for line in md_text.split('\n'):
        line = _clean_xml_text(line)
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
